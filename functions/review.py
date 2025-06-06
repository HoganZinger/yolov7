"""
@author : hogan
time:
function:
"""
import tkinter as tk
from tkinter import DISABLED, NORMAL, StringVar, Label, Button
from tkinter.ttk import Progressbar
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from PIL import Image, ImageTk
from functions.tools import thread_it, select_save, load_image, load_bg
import os
from io import BytesIO
from datetime import datetime

# 用于显示下拉栏
def set_combobox_placeholder(combobox, placeholder):
    combobox.set(placeholder)
    combobox.bind("<FocusIn>", lambda event: combobox_set_normal_style(combobox, placeholder))
    combobox.bind("<<ComboboxSelected>>", lambda event: combobox_set_normal_style(combobox, placeholder))


def combobox_set_normal_style(combobox, placeholder):
    if combobox.get() == placeholder:
        combobox.set('')
        combobox.configure(style='TCombobox')
    elif not combobox.get():
        combobox.set(placeholder)
        combobox.configure(style='Placeholder.TCombobox')


# 更新结果展示列表
def refresh_listbox(listbox, new_items):
    empty_item = ['未检测出有效目标']
    listbox.delete(0, tk.END)  # 清除Listbox中的所有项目
    for item in new_items:
        if not item:
            listbox.insert(tk.END, empty_item)
        else:
            listbox.insert(tk.END, item)  # 添加新的项目到Listbox


# 更新下拉栏内容
def refresh_combobox(combobox, new_items):
    for idx, item in enumerate(new_items):
        if not item:
            new_items[idx] = "未检测出有效目标"
    combobox['values'] = new_items



# 修改选中的检测结果
def change_selection(results_select, option, max):
    selected_index = results_select.current()
    # 未选中任一选项
    if selected_index == -1:
        return
    if option == 'prev':
        if selected_index == 0:
            return
        results_select.current(selected_index - 1)
    if option == 'next':
        if selected_index == max:
            return
        results_select.current(selected_index + 1)
    else:
        return


# 显示给定路径下的图片
def show_image(img_list, window, selected_index):
    image_array = img_list[selected_index]
    image = Image.fromarray(image_array)
    label_width = 300
    label_height = 200
    image = image.resize((label_width, label_height), Image.ANTIALIAS)
    photo = ImageTk.PhotoImage(image)
    label = tk.Label(window, image=photo)
    label.image = photo
    label.place(relx=0.58, rely=0.31)  # 设置图片在窗口中的具体位置

# 获取listbox内容
def get_listbox_items(listbox):
    items = []
    for i in range(listbox.size()):
        items.append(list(listbox.get(i)))
    return items

# 对于选中的条目进行细分展示
def change_list(selected_index, results_list, results_select):
    pre_list = get_listbox_items(results_list)
    # print("pre_list:{}".format(pre_list))
    # print("selected_index:{}".format(selected_index))
    fg_results = pre_list[selected_index]
    print("fg_results:{}".format(fg_results))
    if fg_results == ['未', '检', '测', '出', '有', '效', '目', '标']:
        fg_results = ['未检测出有效目标']
    results_select["values"] = fg_results
    if fg_results[0] != "返回上一级":
        fg_results.insert(0, "返回上一级")
    refresh_listbox(results_list, fg_results)





def result_list_event(event, window, img_list, pre_list, results_list):
    global result_mode
    if result_mode == "picture_selection":
        selection = event.widget.curselection()
        if not selection:
            return
        selected_index = int(selection[0])
        show_image(img_list, window, selected_index)
        change_list(pre_list, selected_index, results_list)
        result_mode = "result_review"
    elif result_mode == "result_review":
        selection = event.widget.curselection()
        if not selection:
            return
        selected_index = int(selection[0])
        if selected_index == 0:
            refresh_listbox(results_list, pre_list)
            result_mode = "picture_selection"
    else:
        return

# 选择结果下拉栏的选中事件，调用结果展示列表的选中事件
def result_select_event(event, window, combobox, listbox):
    selected_index = combobox.current()
    print("Selected")
    if selected_index == -1:
        return
    listbox.selection_clear(0, tk.END)
    print("cleared")
    listbox.selection_set(selected_index)
    print("set")
    result_list_event(event, window)


def generate_report(pre_list, img_list, path_list, button, outdir, window, report_infos, bar):

    button['state'] = DISABLED
    button['text'] = '报告生成中'

    date = report_infos['date_entry'].get()
    serial = report_infos['serial_entry'].get()
    name = report_infos['name_entry'].get()
    item = report_infos['item_entry'].get()

    bar.place(relx=0.28, rely=0.79)
    bar['maximum'] = len(pre_list)
    bar['value'] = 0

    # 处理保存路径
    var = StringVar()
    l = Label(window, textvariable=var, bg="#f0f0f0", font=('宋体', 8), width=20, height=1)
    if outdir == "":
        var.set('保存路径为空')
        l.place(relx=0.63, rely=0.67)
        window.update()
        button['state'] = NORMAL
        button['text'] = '生成检测报告'
        return
    outdir = outdir.replace('\\', '/') + '/'
    if not os.path.exists(outdir) or not os.path.isabs(outdir):  # 检查路径是否存在
        var.set('保存路径不存在')
        l.place(relx=0.63, rely=0.67)
        window.update()
        button['state'] = NORMAL
        button['text'] = '生成检测报告'
        return
    if not os.path.isdir(outdir):  # 检查路径是否是一个文件夹
        var.set('保存路径不是一个文件夹')
        l.place(relx=0.63, rely=0.67)
        window.update()
        button['state'] = NORMAL
        button['text'] = '生成检测报告'
        return
    # 处理编辑信息
    if serial == report_infos['default_serial'] or serial == "":
        var.set('请输入检品编号')
        l.place(relx=0.63, rely=0.67)
        window.update()
        button['state'] = NORMAL
        button['text'] = '生成检测报告'
        print("编号提醒已输出")
        return
    if name == report_infos['default_name'] or name == "":
        var.set('请输入检品名称')
        l.place(relx=0.63, rely=0.67)
        window.update()
        button['state'] = NORMAL
        button['text'] = '生成检测报告'
        print("名字提醒已输出")
        return

    img_name_list = []
    for img_path in enumerate(path_list):
        if isinstance(img_path, tuple):
            path = img_path[1]
            img_path = path
        img_name = os.path.basename(img_path)
        img_name_list.append(img_name)
    # foldername = datetime.datetime.now().strftime('%Y_%m_%d_%H_%M_%S')  # 获取系统时间(年月日时分秒)为文件名
    filename = "编号"+serial
    file_path = f'{outdir}/{filename}.docx'
    # 编辑doc文档
    doc = docx.Document()
    title = doc.add_heading('检验报告', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('日期：{}'.format(date))
    doc.add_paragraph('检品编号：{}'.format(serial))
    doc.add_paragraph('检品名称：{}'.format(name))
    doc.add_paragraph('检验项目：{}'.format(item))
    doc.add_paragraph('检验结果：')
    lenth = len(pre_list)
    # 结果表格
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    headers = ['序号', '图片名称', '结果']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for idx, (img_name, pre_result) in enumerate(zip(img_name_list, pre_list), 1):
        row_cells = table.add_row().cells
        row_cells[0].width = 0.1
        row_cells[0].text = str(idx)
        row_cells[1].text = str(img_name)  # 图片名称

        if pre_result == "未检测出有效目标":
            pre_result = []
        result = f"检测到{len(pre_result)}个目标结构"
        print(idx, img_name, result)
        row_cells[2].text = str(result)  # 检测结果
    doc.add_paragraph('检验结果附图：')

    results_table = doc.add_table(rows=1, cols=3)
    heads = ['序号', '图片', '结果']
    for i, head in enumerate(heads):
        results_table.cell(0, i).text = head
    cnt = 1
    for img_array, list_result, img_name in zip(img_list, pre_list, img_name_list):
        r_cells = results_table.add_row().cells
        r_cells[0].width = docx.shared.Inches(0.2)
        remaining_width = (docx.shared.Inches(6.5) - docx.shared.Inches(0.2)) / 2
        r_cells[0].text = str(cnt)
        r_cells[1].width = remaining_width
        r_cells[2].width = remaining_width

        image = Image.fromarray(img_array)
        image_path = './img_for_report.png'
        image.save(image_path)
        original_size = os.path.getsize(image_path)
        if original_size > 200000:  # 图片过大，需要压缩
            image = Image.open(image_path)
            output = BytesIO()
            image.save(output, format='JPEG', quality=5)
            cell = results_table.cell(cnt, 1)
            cell_paragraph = cell.paragraphs[0]
            run = cell_paragraph.add_run()
            try:
                run.add_picture(image_path, width=Inches(2.4), height=Inches(1.8))
            except Exception as e:
                print(f"Error adding picture. still too big: {e}")
        else:
            cell = results_table.cell(cnt, 1)
            cell_paragraph = cell.paragraphs[0]
            run = cell_paragraph.add_run()
            try:
                run.add_picture(image_path, width=Inches(2.4), height=Inches(1.8))
            except Exception as e:
                print(f"Error adding picture, probably due to wrong path or too big picture: {e}")

        cell = results_table.cell(cnt, 2)
        # cell_paragraph = cell.paragraphs[0]
        if len(list_result) == 0 or list_result == "未检测出有效目标":
            cell = results_table.cell(cnt, 2)
            cell.text = "未检测出有效目标"
        else:
            result_table = cell.add_table(rows=1, cols=2)
            hdr = ['类别', '相似度']
            for i, hd in enumerate(hdr):
                result_table.cell(0, i).text = hd

            for result in list_result:
                print("result: {}".format(result))
                row_cells = result_table.add_row().cells
                row_cells[0].text = result[0]
                row_cells[1].text = result[1]

        cnt += 1
        bar['value'] += 1
        window.update()

    doc.save(file_path)

    button['state'] = NORMAL
    button['text'] = '生成检测报告'
    print("complete!")

    var.set('报告已输出到指定路径')
    l.place(relx=0.55, rely=0.80)
    window.update()
    button['state'] = NORMAL
    button['text'] = '生成检测报告'


# 设置entry的点击事件：在entry为默认文字时清空文字
def on_entry_click(event, entry, default):
    if entry.get() == default:
        entry.delete(0, tk.END)
        entry.config(fg='black')


# 设置entry的失焦事件：在entry空白时设置默认文字
def on_focus_out(event, entry, default):
    if entry.get() == '':
        entry.insert(0, default)
        entry.config(fg='grey')


def show_report_window(window, button, result_dict, pre_list, path_list, img_list):
    button['state'] = DISABLED
    button['text'] = '编辑中'
    report_window = tk.Toplevel(window)
    report_window.title("编辑检测报告")
    report_window.geometry('600x400')
    report_window.resizable(False, False)

    # 设置背景图片
    os_path = os.path.dirname(__file__)
    os_path = os.path.dirname(os_path)  # 返回上级目录
    load_bg(report_window, os_path, 'background/report_window.png')

    save_path_img = load_image(os_path, 'button/report_window/save_path.png', 'small')
    confir_img = load_image(os_path, 'button/report_window/confirm.png', 'small')
    title_img = load_image(os_path, 'background/title/report.png', 'title')
    date_remind_img = load_image(os_path, 'background/remind/date_remind.png', 'small')
    serial_remind_img = load_image(os_path, 'background/remind/serial_remind.png', 'small')
    name_remind_img = load_image(os_path, 'background/remind/name_remind.png', 'small')
    item_remind_img = load_image(os_path, 'background/remind/item_remind.png', 'small')


    # 文字提醒信息
    date_remind = Label(report_window, image=date_remind_img)
    serial_remind = Label(report_window, image=serial_remind_img)
    name_remind = Label(report_window, image=name_remind_img)
    item_remind = Label(report_window, image=item_remind_img)
    title = Label(report_window, image=title_img)
    # 获取检测报告所需的信息
    date_entry = tk.Entry(report_window)
    serial_entry = tk.Entry(report_window)
    name_entry = tk.Entry(report_window)
    item_entry = tk.Entry(report_window)

    # 设置默认值
    default_date = datetime.now().strftime('%Y.%m.%d')
    default_serial = "检品编号"
    default_name = "例：南方菟丝子"
    default_item = "显微鉴别"

    date_entry.insert(0, default_date)
    serial_entry.insert(0, default_serial)
    name_entry.insert(0, default_name)
    item_entry.insert(0, default_item)
    date_entry.config(fg='grey')
    serial_entry.config(fg='grey')
    name_entry.config(fg='grey')
    item_entry.config(fg='grey')

    date_entry.bind('<FocusIn>', lambda event, entry=date_entry, default=default_date: on_entry_click(event, entry, default))
    serial_entry.bind('<FocusIn>', lambda event, entry=serial_entry, default=default_serial: on_entry_click(event, entry, default))
    name_entry.bind('<FocusIn>', lambda event, entry=name_entry, default=default_name: on_entry_click(event, entry, default))
    item_entry.bind('<FocusIn>', lambda event, entry=item_entry, default=default_item: on_entry_click(event, entry, default))

    date_entry.bind('<FocusOut>', lambda event, entry=date_entry, default=default_date: on_focus_out(event, entry, default))
    serial_entry.bind('<FocusOut>', lambda event, entry=serial_entry, default=default_serial: on_focus_out(event, entry, default))
    name_entry.bind('<FocusOut>', lambda event, entry=name_entry, default=default_name: on_focus_out(event, entry, default))
    item_entry.bind('<FocusOut>', lambda event, entry=item_entry, default=default_item: on_focus_out(event, entry, default))

    date_entry.pack()
    serial_entry.pack()
    name_entry.pack()
    item_entry.pack()

    bar = Progressbar(report_window)

    # lable，说明输入框作用
    var2 = StringVar()
    l5 = Label(report_window,
               textvariable=var2,  # 标签的文字
               bg="#f0f0f0",  # 标签背景颜色
               font=('宋体', 8),  # 字体和字体大小
               width=40, height=1  # 标签长宽
               )
    save_folder = Button(report_window, image=save_path_img, width=150, height=40,
                         command=lambda: select_save(var2, l5, save_folder, report_window, result_dict))

    report_infos = {'date_entry': date_entry, 'name_entry': name_entry, 'serial_entry': serial_entry, 'item_entry':
                    item_entry, 'default_serial': default_serial, 'default_name': default_name}
    generate = Button(report_window, image=confir_img, width=150, height=40,
                      command=lambda: generate_report(pre_list, img_list, path_list, generate, result_dict['save_path'],
                                                      report_window, report_infos, bar))

    # 布置按钮
    title.place(relx=0.26, rely=0.05)
    date_remind.place(relx=0.061, rely=0.32)
    date_entry.place(relx=0.25, rely=0.35)
    serial_remind.place(relx=0.05, rely=0.43)
    serial_entry.place(relx=0.25, rely=0.45)
    name_remind.place(relx=0.055, rely=0.53)
    name_entry.place(relx=0.25, rely=0.55)
    item_remind.place(relx=0.05, rely=0.63)
    item_entry.place(relx=0.25, rely=0.662)
    save_folder.place(relx=0.62, rely=0.32)
    generate.place(relx=0.62, rely=0.60)

    button['state'] = NORMAL
    button['text'] = '生成检测报告'

    report_window.mainloop()
